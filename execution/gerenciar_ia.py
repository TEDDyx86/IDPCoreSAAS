import os
import json
import re
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from notificar import notificar_item_ia

load_dotenv()

# Configuração Gemini
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

MODULOS_PATH = ".tmp/modulos_detectados.json"

def extrair_texto_pdf(caminho, max_paginas=50):
    texto = ""
    try:
        doc = fitz.open(caminho)
        # Modo Otimizado: Lemos até 50 páginas para economizar contexto
        for i in range(min(len(doc), max_paginas)):
            texto += f" [PAG {i+1}] " + doc[i].get_text()
        doc.close()
        # Compressão léxica corporativa: Transforma espaços/quebras repetidas em espaço único
        texto = re.sub(r'\s+', ' ', texto).strip()
    except Exception as e:
        print(f"Erro ao ler PDF {caminho}: {e}")
    return texto

def processar_com_gemini(texto, nome_arquivo):
    # Utilizando o motor Flash standard (1.5) que oferece a maior cota do plano Free (1500 por dia)
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    PERSONA: Você é o 'Professor Mentor Ultimate'. Você opera com raciocínio de nível de doutorado. Sua missão é transformar o material abaixo em uma aula de reforço completa, densa e magistral. Você ignora superficialidades e foca no domínio total do assunto pelo aluno.

    TAREFA (Chain of Thought):
    1. Analise TODO o texto fornecido abaixo.
    2. Identifique os conceitos centrais, as nuances técnicas e os exemplos de código/prática.
    3. Construa uma explicação que conecte a teoria com a prática do mercado.
    
    ESTRUTURA DA RESPOSTA (Seções Obrigatórias):
    - CONTEXTO HISTÓRICO E MERCADOLÓGICO: Por que isso existe e como é usado nas Big Techs hoje.
    - NÚCLEO TÉCNICO APROFUNDADO: Explicação exaustiva de conceitos, fórmulas e lógica.
    - LABORATÓRIO E CÓDIGO: Exemplos detalhados de código (Markdown), comandos ou fluxogramas descritos.
    - ANÁLISE CRÍTICA DO MENTOR: O que quase ninguém entende sobre isso, pegadinhas de prova e insights avançados.

    RETORNO: Responda APENAS em JSON:
    {{
        "classe": "SLIDE | ATIVIDADE | ADMIN",
        "resumo_topicos": "Bullet points detalhados dos pontos vitais",
        "resumo_detalhado": "O conteúdo magistral completo, estruturado com os títulos acima. Gere um texto LONGO e rico em detalhes (minímo 1000 palavras se o material permitir).",
        "prazos_detectados": "Datas ou null"
    }}

    DADOS DO ARQUIVO ({nome_arquivo}):
    {texto[:45000]} 
    """
    
    try:
        # Aumentando a temperatura levemente para respostas mais ricas e didáticas
        response = model.generate_content(prompt)
        # Limpa possíveis blocos de código markdown do JSON
        json_text = response.text.replace("```json", "").replace("```", "").strip()
        return json.loads(json_text)
    except Exception as e:
        print(f"Erro na IA PRO para {nome_arquivo}: {e}")
        return None

def salvar_resumos_fisicos(item):
    try:
        resultado = item["ia_processado"]
        disciplina = item["disciplina_nome"].replace("/", "-")
        nome_limpo = item["nome"].replace("/", "-").replace("\\", "-")
        item_id = item["id"]
        
        # Cria estrutura de pastas
        pasta_resumos = os.path.join("resumos", disciplina)
        os.makedirs(pasta_resumos, exist_ok=True)
        
        base_nome = f"{item_id}_{nome_limpo}_Resumo_PRO"
        caminho_txt = os.path.join(pasta_resumos, f"{base_nome}.txt")
        caminho_docx = os.path.join(pasta_resumos, f"{base_nome}.docx")
        
        # 1. Salva TXT
        with open(caminho_txt, "w", encoding="utf-8") as f:
            f.write(f"=== MATERIAL PRO: {item['nome']} ===\n")
            f.write(f"DISCIPLINA: {item['disciplina_nome']}\n")
            f.write(f"CLASSIFICAÇÃO: {resultado['classe']}\n")
            f.write("=" * 40 + "\n\n")
            f.write("TÓPICOS VITAIS:\n")
            f.write(resultado["resumo_topicos"] + "\n\n")
            f.write("AULA MAGISTRAL DO PROFESSOR MENTOR ULTIMATE (PRO):\n")
            f.write(resultado["resumo_detalhado"] + "\n")
            if resultado.get("prazos_detectados"):
                f.write(f"\nPRAZOS E ALERTAS: {resultado['prazos_detectados']}\n")

        # 2. Salva DOCX (Formatado)
        doc = Document()
        doc.add_heading(f"Aula Magistral: {item['nome']}", 0)
        
        p = doc.add_paragraph()
        p.add_run(f"Disciplina: {item['disciplina_nome']}").bold = True
        p.add_run(f"\nClasse Acadêmica: {resultado['classe']}")
        
        doc.add_heading('Pontos Vitais', level=1)
        doc.add_paragraph(resultado["resumo_topicos"])
        
        doc.add_heading('Desenvolvimento Técnico Profundo', level=1)
        # Dividimos por quebras de linha para não ficar uma parede de texto única no Word
        detalhes = resultado["resumo_detalhado"].split('\n')
        for paragrafo in detalhes:
            if paragrafo.strip():
                doc.add_paragraph(paragrafo.strip())
        
        if resultado.get("prazos_detectados"):
            doc.add_heading('Calendário e Prazos', level=2)
            doc.add_paragraph(resultado["prazos_detectados"])
            
        doc.save(caminho_docx)
        print(f"  -> [PRO] Arquivos gerados com sucesso em: {pasta_resumos}")
        return True
    except Exception as e:
        print(f"Erro ao salvar arquivos PRO: {e}")
        return False

def main():
    if not os.path.exists(MODULOS_PATH):
        return

    with open(MODULOS_PATH, "r", encoding="utf-8") as f:
        materiais = json.load(f)

    for item in materiais:
        # Se você tem o PRO, queremos re-processar para ter a qualidade máxima
        # ou apenas processar os que ainda não têm o sufixo PRO?
        # Para evitar custo infinito, vamos processar apenas os novos ou se o usuário limpar.
        if "caminho_local" in item and "ia_processado" not in item:
            print(f"Iniciando Processamento PRO para: {item['nome']}...")
            
            texto = extrair_texto_pdf(item["caminho_local"])
            if not texto:
                continue
                
            resultado = processar_com_gemini(texto, item["nome"])
            if resultado:
                item["ia_processado"] = resultado
                print(f"  -> Classe PRO detectada: {resultado['classe']}")
                
                # Dispara notificação imediata
                notificar_item_ia(item)
                
                # Gera arquivos físicos (.txt e .docx)
                salvar_resumos_fisicos(item)
                
                # Atualiza após cada arquivo
                with open(MODULOS_PATH, "w", encoding="utf-8") as f:
                    json.dump(materiais, f, indent=2, ensure_ascii=False)

if __name__ == "__main__":
    main()
